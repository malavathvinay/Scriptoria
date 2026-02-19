import os
from flask import Flask, render_template, request, jsonify, session, send_file
from groq import Groq
from dotenv import load_dotenv
import io
import uuid
import base64

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors

# python-docx imports
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

# Load .env — override=True ensures a running server picks up key changes immediately
load_dotenv(override=True)

app = Flask(__name__)
app.secret_key = 'scriptoria_secret_key_2024'

# Server-side store — avoids Flask's 4 KB cookie limit for large AI outputs
_results_store: dict = {}

GROQ_MODEL  = "llama-3.3-70b-versatile"
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# —— Startup: confirm HF key is loaded ——
_hf_key_preview = os.getenv("HF_API_KEY", "")
if _hf_key_preview and _hf_key_preview != "hf_your_token_here":
    print(f"[Scriptoria] ✅  HF_API_KEY loaded: {_hf_key_preview[:8]}...")
else:
    print("[Scriptoria] ⚠️  HF_API_KEY not set — image generation will prompt for setup.")


# Hugging Face client singleton (re-reads env every call so restarts aren't needed)
_hf_client = None

def get_hf_client():
    global _hf_client
    from huggingface_hub import InferenceClient
    # Always re-read from env — load_dotenv(override=True) keeps it fresh
    api_key = os.getenv("HF_API_KEY", "").strip()
    if not api_key or api_key == "hf_your_token_here":
        _hf_client = None       # reset so next attempt retries
        raise RuntimeError("HF_API_KEY_NOT_CONFIGURED")
    if _hf_client is None:      # only create once per valid key
        _hf_client = InferenceClient(
            provider="hf-inference",
            api_key=api_key,
        )
    return _hf_client

# ---------------------------------------------------------------------------
# Groq helper
# ---------------------------------------------------------------------------

def generate_with_groq(prompt):
    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=GROQ_MODEL,
            temperature=0.7,
            max_tokens=4096,
        )
        return chat_completion.choices[0].message.content.strip()
    except Exception as e:
        return f"[ERROR] Groq API error: {str(e)}"

# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

def build_screenplay_prompt(story):
    return f"""You are a professional Hollywood screenwriter. Write a complete, industry-standard screenplay based on the following story idea.

STRICT FORMATTING RULES:
- Scene headings MUST be in ALL CAPS (e.g., INT. COFFEE SHOP - DAY)
- Character names before dialogue MUST be in ALL CAPS and centered
- Action lines in present tense, concise
- Include at least 3 scenes with dialogue
- Use proper screenplay structure: Setup, Confrontation, Resolution

STORY IDEA:
{story}

Write the full screenplay now:"""


def build_characters_prompt(story):
    return f"""You are a professional character designer and casting director. Based on the story below, create detailed character profiles for all major and supporting characters.

For EACH character provide:
- NAME (in ALL CAPS)
- AGE & PHYSICAL DESCRIPTION
- PERSONALITY TRAITS (3-5 bullet points)
- BACKSTORY (2-3 sentences)
- ROLE IN STORY
- SUGGESTED CASTING TYPE (e.g., "Mid-30s rugged male lead")
- CHARACTER ARC (how they change by the end)

STORY:
{story}

Generate all character profiles now:"""


def build_sound_design_prompt(story):
    return f"""You are an award-winning sound designer and composer. Create a comprehensive sound design plan for the following story.

Include:
1. OVERALL SONIC PALETTE (describe the emotional tone and audio world)
2. SCENE-BY-SCENE SOUND MAP:
   - Ambient sounds / room tone
   - Key sound effects (SFX)
   - Music cues (genre, tempo, instrumentation)
3. LEITMOTIFS: Recurring musical themes for main characters/concepts
4. SILENCE USAGE: Where silence is used for dramatic effect
5. TECHNICAL NOTES: Recommended recording techniques or sound libraries

STORY:
{story}

Create the full sound design plan now:"""


def build_script_breakdown_prompt(story):
    return f"""You are an experienced film production coordinator. Create a detailed script breakdown from the following story.

For EACH SCENE provide a breakdown table with:
- SCENE NUMBER & HEADING (INT/EXT, LOCATION, TIME)
- CAST REQUIREMENTS: Named characters appearing in scene
- PROPS: All physical objects needed
- LOCATIONS: Specific set or location requirements
- WARDROBE: Key costume notes
- SPECIAL REQUIREMENTS: Stunts, VFX, special equipment
- ESTIMATED SHOOT TIME: (in hours)

Also include at the end:
- TOTAL CAST LIST (all unique characters)
- MASTER PROPS LIST (all unique props)
- UNIQUE LOCATIONS LIST

STORY:
{story}

Generate the complete script breakdown now:"""


def build_shot_list_prompt(story):
    return f"""You are a cinematographer and director of photography (DP). Create a professional shot list for the following story.

For EACH SHOT provide:
- SHOT NUMBER
- SCENE REFERENCE
- SHOT TYPE: (e.g., ECU - Extreme Close Up, CU, MS, WS, EWS)
- CAMERA ANGLE: (e.g., Low Angle, High Angle, Bird's Eye, Dutch Tilt, Eye Level)
- CAMERA MOVEMENT: (e.g., Static, Pan, Tilt, Dolly, Handheld, Steadicam, Crane)
- LENS SUGGESTION: (e.g., 24mm wide, 35mm standard, 50mm normal, 85mm portrait, 135mm telephoto)
- LIGHTING MOOD: (e.g., Noir - high contrast, High-key - bright/flat, Golden Hour, Practical only, Motivated)
- DESCRIPTION: What the shot shows and its emotional purpose
- ESTIMATED DURATION: (in seconds)

STORY:
{story}

Generate the complete professional shot list now:"""


def build_image_prompt_from_shot(shot_description):
    return f"""You are a professional Hollywood cinematographer and AI prompt engineer.

Your task: Convert the following shot description into a highly detailed, cinematic AI image generation prompt.

Include ALL of these elements:
- Shot type (close-up, wide shot, aerial, extreme close-up, over-the-shoulder, etc.)
- Subject details (age, appearance, clothing, expression, body language)
- Action happening in the frame
- Environment/location with rich atmospheric detail
- Time of day (golden hour, blue hour, midnight, noon, etc.)
- Lighting style (Rembrandt, softbox, motivated practical, neon, candlelight, etc.)
- Camera angle (low angle, bird's eye, dutch tilt, eye level, etc.)
- Lens type and effect (24mm wide-angle, 35mm, 50mm standard, 85mm portrait bokeh, 135mm telephoto compression)
- Mood and emotion (tense, melancholic, euphoric, foreboding, romantic, etc.)
- Cinematic style reference (e.g., Blade Runner 2049, No Country for Old Men, La La Land, etc.)
- Color grading style (desaturated teal-orange, warm analog, cool monochrome, vivid Kodak)

End the prompt with: ultra realistic, 4K, high detail, film still, cinematic photography, anamorphic lens flare

SHOT DESCRIPTION:
{shot_description}

Output ONLY the final image prompt. No explanation, no preamble, no labels. Just the prompt itself."""

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/set_user', methods=['POST'])
def set_user():
    data = request.get_json()
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': 'Name is required'}), 400
    session['user_name'] = name
    if 'sid' not in session:
        session['sid'] = str(uuid.uuid4())
    return jsonify({'success': True, 'name': name})


@app.route('/get_user', methods=['GET'])
def get_user():
    return jsonify({'name': session.get('user_name', '')})


@app.route('/generate', methods=['POST'])
def generate():
    data  = request.get_json()
    story = data.get('story', '').strip()

    if not story:
        return jsonify({'error': 'Story input is required'}), 400

    if 'sid' not in session:
        session['sid'] = str(uuid.uuid4())
    sid = session['sid']

    session['story'] = story[:200]  # keep cookie small

    prompts = {
        'screenplay':       build_screenplay_prompt(story),
        'characters':       build_characters_prompt(story),
        'sound_design':     build_sound_design_prompt(story),
        'script_breakdown': build_script_breakdown_prompt(story),
        'shot_list':        build_shot_list_prompt(story),
    }

    results = {key: generate_with_groq(prompt) for key, prompt in prompts.items()}

    _results_store[sid] = results
    return jsonify({'success': True, 'results': results})


@app.route('/get_results', methods=['GET'])
def get_results():
    sid = session.get('sid')
    return jsonify(_results_store.get(sid, {}) if sid else {})


@app.route('/generate_shot_image', methods=['POST'])
def generate_shot_image():
    """Convert a shot description → cinematic AI prompt → HF SDXL image (base64)."""
    data = request.get_json(force=True, silent=True) or {}
    shot_text = data.get('shot_description', '').strip()

    if not shot_text:
        return jsonify({'error': 'shot_description is required'}), 400

    # Step 1 — Groq crafts a rich Hollywood-DP image prompt
    image_prompt = generate_with_groq(build_image_prompt_from_shot(shot_text))
    if image_prompt.startswith('[ERROR]'):
        return jsonify({'error': image_prompt}), 500

    # Step 2 — Call Hugging Face SDXL (free tier, returns a PIL Image)
    try:
        client = get_hf_client()
        pil_image = client.text_to_image(
            prompt=image_prompt[:900],          # keep within model token limits
            model="stabilityai/stable-diffusion-xl-base-1.0",
        )
    except RuntimeError as env_err:
        if "HF_API_KEY_NOT_CONFIGURED" in str(env_err):
            return jsonify({
                'setup_required': True,
                'steps': [
                    'Visit https://huggingface.co/settings/tokens',
                    'Click "New token" → choose Role: Read → click Create',
                    'Copy the token (it starts with  hf_...)',
                    'Open the file  .env  in your project folder',
                    'Replace  hf_your_token_here  with your real token',
                    'Save the file, then restart the server:  python app.py',
                ]
            }), 200
        return jsonify({'error': str(env_err)}), 500
    except Exception as hf_err:
        return jsonify({'error': f'HF image generation failed: {str(hf_err)}'}), 500

    # Step 3 — Convert PIL Image → base64 PNG data-URI (no external host needed)
    buf = io.BytesIO()
    pil_image.save(buf, format='PNG')
    buf.seek(0)
    b64 = base64.b64encode(buf.read()).decode('utf-8')
    data_uri = f"data:image/png;base64,{b64}"

    return jsonify({
        'success': True,
        'image_prompt': image_prompt,
        'image_url': data_uri,   # frontend <img src> accepts data-URIs directly
    })

# ---------------------------------------------------------------------------
# Export routes
# ---------------------------------------------------------------------------

@app.route('/export/<content_type>/<fmt>', methods=['POST'])
def export(content_type, fmt):
    data    = request.get_json(force=True, silent=True) or {}
    content = data.get('content', '').strip()
    title   = content_type.replace('_', ' ').title()

    # Fallback: server-side store
    if not content:
        sid     = session.get('sid')
        results = _results_store.get(sid, {}) if sid else {}
        content = results.get(content_type, '')

    if not content:
        return jsonify({'error': 'No content found. Please generate first.'}), 404

    filename = f"scriptoria_{content_type}"

    # ---- TXT ----------------------------------------------------------------
    if fmt == 'txt':
        buf = io.BytesIO()
        buf.write(f"SCRIPTORIA — {title.upper()}\n".encode('utf-8'))
        buf.write(("=" * 60 + "\n\n").encode('utf-8'))
        buf.write(content.encode('utf-8'))
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                         as_attachment=True, download_name=f"{filename}.txt")

    # ---- PDF ----------------------------------------------------------------
    elif fmt == 'pdf':
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'ScriptoriaTitle', parent=styles['Title'],
            fontName='Times-Bold', fontSize=18,
            textColor=colors.HexColor('#1a1a2e'),
            spaceAfter=6, alignment=TA_CENTER
        )
        subtitle_style = ParagraphStyle(
            'ScriptoriaSubtitle', parent=styles['Normal'],
            fontName='Times-Italic', fontSize=11,
            textColor=colors.HexColor('#555555'),
            spaceAfter=20, alignment=TA_CENTER
        )
        body_style = ParagraphStyle(
            'ScriptoriaBody', parent=styles['Normal'],
            fontName='Times-Roman', fontSize=11,
            leading=16.5, spaceAfter=8, alignment=TA_LEFT
        )

        elements = [
            Paragraph("SCRIPTORIA", title_style),
            Paragraph(title, subtitle_style),
            Spacer(1, 0.2 * inch),
        ]
        for para in content.split('\n'):
            para = para.strip()
            if para:
                para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(para, body_style))
            else:
                elements.append(Spacer(1, 0.1 * inch))

        doc.build(elements)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name=f"{filename}.pdf")

    # ---- DOCX ---------------------------------------------------------------
    elif fmt == 'docx':
        buf = io.BytesIO()
        doc = Document()

        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        heading = doc.add_heading('SCRIPTORIA', level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Cambria'
            run.font.size = Pt(20)

        sub = doc.add_heading(title, level=1)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.name = 'Cambria'
            run.font.size = Pt(14)

        doc.add_paragraph()

        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Cambria'
                run.font.size = Pt(11)
            pPr     = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)

        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{filename}.docx"
        )

    return jsonify({'error': 'Invalid format. Use txt, pdf, or docx'}), 400

# ---------------------------------------------------------------------------

if __name__ == '__main__':
    app.run(debug=True, port=5000)
